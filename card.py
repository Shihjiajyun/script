"""
手卡生成器 - script_to_cards.py
====================================
用法：
    python script_to_cards.py 講稿.md
    python script_to_cards.py 講稿.md -o 輸出手卡.docx
    python script_to_cards.py 講稿.md --max-chars 140

MD 格式規則：
    ## 章節標題        → 藍色小標（顯示在格子頂端）
    姓名：台詞內容     → 對話行（姓名後接全形或半形冒號）
    ★                 → 提示符號，印出為紅色★
    空行               → 忽略（僅作分隔）
    純文字行（無冒號） → 接在上一位司儀名字下方（延續台詞）

排版邏輯：
    - 每格最多 MAX_CHARS 個字（預設 140，含姓名）
    - 同一位司儀的台詞不拆格
    - 如果目前格剩餘空間放不下下一段台詞，就換下一格
    - 每頁 4 列（共 8 格），滿了自動往下一頁

安裝依賴：
    pip install python-docx
"""

import re
import sys
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── 版面常數（對應手卡模板）────────────────────────────────────
PAGE_W_CM   = 20.99   # A4 寬
PAGE_H_CM   = 29.69   # A4 高
MARGIN_CM   = 1.27    # 上下左右邊距（約 720 DXA）
CELL_W_CM   = 9.26    # 每格寬（5220 DXA ≈ 9.26cm）
ROW_H_CM    = 5.78    # 每列高（3261 DXA ≈ 5.78cm）
FONT_NAME   = '標楷體'
FONT_SIZE   = 11      # pt
LINE_SPACING = 1.15   # 行距倍數

SECTION_COLOR = RGBColor(0x2E, 0x74, 0xB5)   # 章節標題藍色
STAR_COLOR    = RGBColor(0xCC, 0x00, 0x00)    # ★ 紅色
NAME_COLOR    = RGBColor(0x00, 0x00, 0x00)    # 姓名黑色

# ── 資料結構 ────────────────────────────────────────────────────

class Block:
    """一段台詞（一位司儀說的一整段話）"""
    def __init__(self, speaker: str, lines: list[str], section: str = ''):
        self.speaker = speaker      # 姓名，如「孟晴」
        self.lines   = lines        # 台詞行列表（可多行）
        self.section = section      # 所屬章節標題（只在章節第一個 block 有值）

    @property
    def char_count(self) -> int:
        """計算字數（姓名＋冒號＋全部台詞）"""
        text = self.speaker + '：' + ''.join(self.lines)
        return len(text)

    def __repr__(self):
        preview = ''.join(self.lines)[:20]
        return f'Block({self.speaker!r}, {self.char_count}字, {preview!r}...)'


# ── MD 解析 ─────────────────────────────────────────────────────

def parse_md(md_text: str) -> list[Block]:
    """
    解析 MD 格式講稿，回傳 Block 列表。
    格式：
        ## 章節標題
        姓名：台詞（全形或半形冒號）
        續行（無冒號開頭，接在上一個 block）
    """
    blocks: list[Block] = []
    current_section = ''
    current_speaker = ''
    current_lines: list[str] = []
    pending_section = ''   # 等待放入下一個 block

    def flush():
        nonlocal current_speaker, current_lines
        if current_speaker and current_lines:
            b = Block(current_speaker, list(current_lines), pending_section if not blocks else '')
            # 章節標題放在該章節第一個 block
            blocks.append(b)
        current_speaker = ''
        current_lines = []

    lines = md_text.splitlines()
    section_for_next = ''

    for raw in lines:
        line = raw.strip()

        # 章節標題
        if line.startswith('## '):
            flush()
            section_for_next = line[3:].strip()
            continue

        # 空行：忽略
        if not line:
            continue

        # 對話行（姓名：台詞）
        m = re.match(r'^(.{1,6})[：:]\s*(.*)', line)
        if m:
            flush()
            current_speaker = m.group(1).strip()
            tail = m.group(2).strip()
            current_lines = [tail] if tail else []
            # 附上章節標題（只給這批第一個 block）
            if section_for_next:
                # 標記到下一個被 flush 出去的 block
                pass
        else:
            # 續行（無冒號）
            if current_speaker:
                current_lines.append(line)
            # 若還沒有 speaker，跳過

    flush()

    # 補上章節標題：重新掃一次，把 ## 對應到後面第一個 block
    # 重新解析一次以確保章節對應正確
    return _parse_md_with_sections(md_text)


def _parse_md_with_sections(md_text: str) -> list[Block]:
    blocks: list[Block] = []
    current_section = ''
    next_section = ''
    current_speaker = ''
    current_lines: list[str] = []
    section_assigned = False

    def flush():
        nonlocal current_speaker, current_lines, next_section, section_assigned
        if current_speaker and (current_lines or True):
            content = [l for l in current_lines if l]  # 保留空字串表示空行
            if not content and not current_speaker:
                return
            sec = next_section if not section_assigned else ''
            # 實際上我們讓每個章節第一個block帶 section
            b = Block(current_speaker, list(current_lines), sec)
            blocks.append(b)
            next_section = ''
            section_assigned = True
        current_speaker = ''
        current_lines = []

    lines = md_text.splitlines()

    for raw in lines:
        line = raw.strip()

        if line.startswith('## '):
            flush()
            next_section = line[3:].strip()
            section_assigned = False
            continue

        if not line:
            continue

        m = re.match(r'^(.{1,6})[：:]\s*(.*)', line)
        if m:
            flush()
            current_speaker = m.group(1).strip()
            tail = m.group(2).strip()
            current_lines = [tail] if tail else []
        else:
            if current_speaker:
                current_lines.append(line)

    flush()
    return blocks


# ── 排版邏輯 ────────────────────────────────────────────────────

def layout_blocks(blocks: list[Block], max_chars: int = 120) -> list[list[Block | None]]:
    """
    將 blocks 排入格子。
    回傳 list of cells，每個 cell 是 list[Block]（同一格的 blocks）。
    None 代表空格（填充）。

    規則：
    1. 每格最多 max_chars 字（含姓名）
    2. 同一位司儀的台詞不拆格
    3. 放不下就換下一格
    """
    cells: list[list[Block]] = []
    current_cell: list[Block] = []
    current_chars = 0

    for block in blocks:
        blen = block.char_count
        # 如果單一 block 超過 max_chars，直接放一格（不強拆）
        fits = (current_chars + blen) <= max_chars

        if fits and current_chars > 0:
            current_cell.append(block)
            current_chars += blen
        else:
            # 換格
            if current_cell:
                cells.append(current_cell)
            current_cell = [block]
            current_chars = blen

    if current_cell:
        cells.append(current_cell)

    return cells


# ── DOCX 生成 ────────────────────────────────────────────────────

def set_cell_height(row, height_cm: float):
    """設定列的精確高度"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    twips = int(height_cm * 567)  # 1cm = 567 twips (DXA)
    trHeight.set(qn('w:val'), str(twips))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)


def set_cell_border(cell, sides: dict):
    """設定儲存格邊框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    border_map = {'top': 'top', 'bottom': 'bottom', 'left': 'left', 'right': 'right'}
    for side, show in sides.items():
        b = OxmlElement(f'w:{side}')
        if show:
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '8')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), 'auto')
        else:
            b.set(qn('w:val'), 'none')
            b.set(qn('w:sz'), '0')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), 'auto')
        tcBorders.append(b)
    tcPr.append(tcBorders)


def set_cell_margin(cell, top=57, bottom=57, left=120, right=120):
    """設定儲存格內距（DXA）"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def set_para_spacing(para, line_spacing: float = LINE_SPACING):
    """設定段落行距"""
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), str(int(240 * line_spacing)))
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)


def add_run_with_stars(para, text: str, bold: bool = False,
                       color: RGBColor = None, font: str = FONT_NAME, size: int = FONT_SIZE):
    """將文字中的 ★ 拆開，★ 用紅色粗體，其餘照常"""
    parts = text.split('★')
    for i, part in enumerate(parts):
        if i > 0:
            run = para.add_run('★')
            run.bold = True
            run.font.color.rgb = STAR_COLOR
            run.font.name = font
            run.font.size = Pt(size)
            run._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), font)
        if part:
            run = para.add_run(part)
            run.bold = bold
            run.font.name = font
            run.font.size = Pt(size)
            run._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), font)
            if color:
                run.font.color.rgb = color


def fill_cell(cell, cell_blocks: list[Block], cell_number: int = None):
    """把 blocks 的內容寫入儲存格"""
    # 清除預設空段落
    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)

    # 章節標題（只有第一個 block 可能帶有）
    first_section = next((b.section for b in cell_blocks if b.section), '')
    if first_section:
        p = cell.add_paragraph()
        run = p.add_run(f'【{first_section}】')
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE)
        run.font.color.rgb = SECTION_COLOR
        run._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), FONT_NAME)
        set_para_spacing(p)

    for block in cell_blocks:
        # 第一行：姓名 + 首行台詞
        first_line = block.lines[0] if block.lines else ''
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # 設定懸掛縮排（姓名後的換行縮進）
        pPr = p._p.get_or_add_pPr()
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '660')
        ind.set(qn('w:hanging'), '660')
        pPr.append(ind)
        set_para_spacing(p)

        # 姓名（粗體）
        name_run = p.add_run(block.speaker + '：')
        name_run.bold = True
        name_run.font.name = FONT_NAME
        name_run.font.size = Pt(FONT_SIZE)
        name_run._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), FONT_NAME)

        # 首行台詞
        add_run_with_stars(p, first_line)

        # 後續行
        for extra_line in block.lines[1:]:
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pPr2 = p2._p.get_or_add_pPr()
            ind2 = OxmlElement('w:ind')
            ind2.set(qn('w:left'), '660')
            pPr2.append(ind2)
            set_para_spacing(p2)
            add_run_with_stars(p2, extra_line)

    if cell_number is not None:
        p_num = cell.add_paragraph()
        p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_num = p_num.add_run(str(cell_number))
        run_num.font.name = FONT_NAME
        run_num.font.size = Pt(9)
        run_num.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run_num._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), FONT_NAME)


def generate_docx(cells: list[list[Block]], output_path: str, show_pagenum: bool = False):
    """生成手卡 DOCX"""
    doc = Document()

    # 頁面設定
    section = doc.sections[0]
    section.page_width  = Cm(PAGE_W_CM)
    section.page_height = Cm(PAGE_H_CM)
    section.top_margin    = Cm(MARGIN_CM)
    section.bottom_margin = Cm(MARGIN_CM)
    section.left_margin   = Cm(MARGIN_CM)
    section.right_margin  = Cm(MARGIN_CM)

    # 每頁 4 列（8 格），超過自動換頁（Word 會自動處理）
    # 每頁塞 4 列的 table，滿了插 page break 再開新 table

    ROWS_PER_PAGE = 4
    total_cells = len(cells)

    # 補齊格數為偶數（左右配對）
    if total_cells % 2 != 0:
        cells.append([])

    # 把 cells 配對成 rows（每兩格一列：左、右）
    rows_data = [(cells[i], cells[i+1]) for i in range(0, len(cells), 2)]

    row_idx = 0
    while row_idx < len(rows_data):
        # 取這一頁的列
        page_rows = rows_data[row_idx: row_idx + ROWS_PER_PAGE]
        row_idx += ROWS_PER_PAGE

        # 建立表格
        table = doc.add_table(rows=len(page_rows), cols=2)
        table.style = 'Table Grid'

        # 設定整體表格寬度
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), '10440')  # 5220 * 2 DXA
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)

        # 設定欄寬
        tblGrid = OxmlElement('w:tblGrid')
        for _ in range(2):
            gridCol = OxmlElement('w:gridCol')
            gridCol.set(qn('w:w'), '5220')
            tblGrid.append(gridCol)
        tbl.insert(0, tblGrid)

        for r_idx, (left_blocks, right_blocks) in enumerate(page_rows):
            row = table.rows[r_idx]
            set_cell_height(row, ROW_H_CM)

            left_cell  = row.cells[0]
            right_cell = row.cells[1]

            # 設定欄寬
            for cell, w in [(left_cell, 5220), (right_cell, 5220)]:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(w))
                tcW.set(qn('w:type'), 'dxa')
                tcPr.append(tcW)

            # 邊框：只有上下有框線，左右無框線
            set_cell_border(left_cell,  {'top': True,  'bottom': True,  'left': False, 'right': False})
            set_cell_border(right_cell, {'top': True,  'bottom': True,  'left': False, 'right': False})

            # 內距
            set_cell_margin(left_cell)
            set_cell_margin(right_cell)

            # 計算格號（全域格子編號，從 1 起算）
            global_row = row_idx - len(page_rows) + r_idx
            left_num  = global_row * 2 + 1 if show_pagenum else None
            right_num = global_row * 2 + 2 if show_pagenum else None

            # 填入內容
            if left_blocks:
                fill_cell(left_cell, left_blocks, left_num)
            if right_blocks:
                fill_cell(right_cell, right_blocks, right_num)

        # 如果還有下一頁，插入分頁符號
        if row_idx < len(rows_data):
            doc.add_page_break()

    doc.save(output_path)
    print(f'✅ 已輸出：{output_path}')
    print(f'   共 {len(cells)} 格，{len(rows_data)} 列')


# ── 主程式 ───────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description='將 MD 格式講稿轉換為手卡 DOCX',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__
    )
    parser.add_argument('input', help='輸入的 MD 講稿檔案路徑')
    parser.add_argument('-o', '--output', help='輸出的 DOCX 路徑（預設：與輸入同名）')
    parser.add_argument('--max-chars', type=int, default=140,
                        help='每格最多字數（預設：140）')
    parser.add_argument('--debug', action='store_true',
                        help='印出排版結果但不輸出 DOCX')
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f'❌ 找不到檔案：{args.input}')
        sys.exit(1)

    output_path = args.output or input_path.with_suffix('.docx')

    md_text = input_path.read_text(encoding='utf-8')
    blocks = parse_md(md_text)

    print(f'📄 解析完成：共 {len(blocks)} 段台詞')

    if args.debug:
        cells = layout_blocks(blocks, max_chars=args.max_chars)
        for i, cell in enumerate(cells):
            total = sum(b.char_count for b in cell)
            print(f'\n── 格 {i+1}（{total} 字）──')
            for b in cell:
                section_tag = f'[{b.section}] ' if b.section else ''
                print(f'  {section_tag}{b.speaker}：{"".join(b.lines)[:40]}...')
        return

    cells = layout_blocks(blocks, max_chars=args.max_chars)
    print(f'📐 排版完成：共 {len(cells)} 格，最大字數限制 {args.max_chars}')

    generate_docx(cells, str(output_path))


if __name__ == '__main__':
    main()